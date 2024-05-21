from docx import Document
from docx.shared import Cm
import matplotlib.pyplot as plt
import seaborn as sns
from typing import Tuple
from matplotlib.figure import Axes
from general_classes import Subject,Experiment
import os
sns.set_theme()


class Report():
    """
    This class generates a report for a subject that has participated in experiments.
    
    
    Attributes:
    subj (Subject): The subject for whom the report is generated.
    exp (Experiment): The experiment data used in the report.
    img_path (str): Path to save the generated images for the report.
    """

    def __init__(self, subject: Subject, experiment: Experiment, 
    img_path = 'Images/tmp.png'):
        self.subj = subject
        self.exp = experiment
        self.img_path = img_path

    def create_graphics(self) -> Tuple[Axes, Axes]:
        """
        Generates a line plot and a histogram for the experiment data.
        
        
        Returns:
        Tuple[Axes, Axes]: A tuple containing the Axes objects for the line plot and histogram.
        """
        fig, ax = plt.subplots(1,2, figsize=(11,5))
        ax[0].plot(self.exp.data)
        ax[0].set_ylabel('Reaction Time (s)')
        ax[0].set_xlabel('Trial number')


        sns.histplot(self.exp.data, ax=ax[1], kde=True, stat="density",
kde_kws=dict(cut=3))
        ax[1].set_ylabel('Frequencia')
        ax[1].set_xlabel('Reaction Time (seconds)')
        return ax

    def generate_report(self) -> str:
        """
        Generates a Word document report based on the subject and experiment data.
        
        
        Returns:
        str: The file path of the generated Word document.
        
        
        Note:
        The report includes subject information, experiment results, and graphical data representations.            
        """
  # Instantiate a Document
        document = Document()
    # Add a title, 0 is for the hyerarchy of the title.
        document.add_heading(f'Informe de Resultados', 0)

        # Subject information in a new paragraph
        p = document.add_paragraph(f'Subject ID: ')
        p.add_run(f'{self.subj.subject_id}').bold = True
        # List the subject information
        for k, v in self.subj.info.items():
            document.add_paragraph(f'{k}= {v}',
                                   style='List Bullet')

        # Experiment results (you can check the different styles)
        p = document.add_paragraph(f'Resultados de {self.exp.exp_name}',
                                   style='Intense Quote')
        # List the experiment results
        for k, v in self.subj.results.items():
            document.add_paragraph(f'{k}= {v}',
                                   style='List Bullet')

        # Add graphs
        ax = self.create_graphics()

        # Create Images folder if not exists
        directory = 'Images'
        if not os.path.exists(directory):
            os.makedirs(directory)


        plt.savefig(self.img_path) # Temporary path for the image
        plt.close() # Close the plot to only save it
        document.add_picture(self.img_path, width=Cm(14))


        # Save the Word document
        fp = f'reporte_{self.subj.subject_id}.docx'
        document.save(fp)
        return fp